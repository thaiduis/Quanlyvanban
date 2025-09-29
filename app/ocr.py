from __future__ import annotations

import os
import re
import cv2
import numpy as np
from typing import List, Tuple, Optional
import logging

import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import pdfplumber
from pdf2image import convert_from_path
from docx import Document
import zipfile
import xml.etree.ElementTree as ET
import docx2txt

from dotenv import load_dotenv

load_dotenv()

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

TESSERACT_CMD = os.getenv("TESSERACT_CMD")
if TESSERACT_CMD and os.path.exists(TESSERACT_CMD):
	pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

POPPLER_BIN = os.getenv("POPPLER_BIN")

# Vietnamese text correction patterns
VIETNAMESE_CORRECTIONS = {
    # Common OCR mistakes for Vietnamese
    '0': 'o', '1': 'l', '5': 's', '8': 'B',
    'rn': 'm', 'cl': 'd', 'ci': 'a',
    'I': 'l', 'O': 'o', 'S': 's',
    # Vietnamese specific corrections
    'a': 'à', 'e': 'è', 'i': 'ì', 'o': 'ò', 'u': 'ù',
    'an': 'ân', 'en': 'ên', 'in': 'în', 'on': 'ôn', 'un': 'ûn',
    'ang': 'âng', 'eng': 'êng', 'ing': 'îng', 'ong': 'ông', 'ung': 'ûng'
}

# Extended Vietnamese words for context correction
VIETNAMESE_WORDS = [
    # Common words
    'và', 'của', 'trong', 'với', 'được', 'không', 'có', 'là', 'để', 'từ',
    'này', 'đó', 'đã', 'sẽ', 'đang', 'đã', 'được', 'cần', 'phải', 'nên',
    'tại', 'về', 'theo', 'như', 'khi', 'nếu', 'mặc', 'dù', 'tuy', 'nhiên',
    
    # Academic and formal words
    'nghiên', 'cứu', 'phân', 'tích', 'đánh', 'giá', 'kết', 'quả', 'phương',
    'pháp', 'thực', 'nghiệm', 'lý', 'thuyết', 'ứng', 'dụng', 'phát', 'triển',
    'công', 'nghệ', 'thông', 'tin', 'hệ', 'thống', 'quản', 'lý', 'dữ', 'liệu',
    
    # Business and professional words
    'doanh', 'nghiệp', 'công', 'ty', 'tổ', 'chức', 'quản', 'trị', 'kinh', 'doanh',
    'thị', 'trường', 'khách', 'hàng', 'sản', 'phẩm', 'dịch', 'vụ', 'chi', 'phí',
    'lợi', 'nhuận', 'đầu', 'tư', 'tài', 'chính', 'ngân', 'hàng', 'bảo', 'hiểm',
    
    # Technical terms
    'máy', 'tính', 'phần', 'mềm', 'ứng', 'dụng', 'website', 'internet', 'mạng',
    'lưới', 'dữ', 'liệu', 'cơ', 'sở', 'dữ', 'liệu', 'thuật', 'toán', 'lập', 'trình',
    'phát', 'triển', 'web', 'mobile', 'android', 'ios', 'server', 'client',
    
    # Education words
    'học', 'sinh', 'sinh', 'viên', 'giáo', 'viên', 'trường', 'đại', 'học', 'cao',
    'đẳng', 'trung', 'học', 'tiểu', 'học', 'mầm', 'non', 'giáo', 'dục', 'đào', 'tạo',
    'nghiên', 'cứu', 'khoa', 'học', 'luận', 'án', 'luận', 'văn', 'báo', 'cáo',
    
    # Government and legal words
    'chính', 'phủ', 'nhà', 'nước', 'luật', 'pháp', 'quy', 'định', 'nghị', 'định',
    'thông', 'tư', 'chỉ', 'thị', 'quyết', 'định', 'bộ', 'trưởng', 'thủ', 'tướng',
    'tổng', 'thống', 'quốc', 'hội', 'hội', 'đồng', 'nhân', 'dân', 'tỉnh', 'thành',
    
    # Medical and health words
    'bác', 'sĩ', 'y', 'tá', 'bệnh', 'viện', 'phòng', 'khám', 'thuốc', 'chữa', 'bệnh',
    'sức', 'khỏe', 'y', 'tế', 'công', 'cộng', 'dịch', 'bệnh', 'virus', 'vi', 'khuẩn',
    
    # Common verbs
    'làm', 'làm', 'việc', 'học', 'học', 'tập', 'đi', 'đến', 'về', 'ở', 'sống',
    'ăn', 'uống', 'ngủ', 'nghỉ', 'chơi', 'xem', 'nghe', 'đọc', 'viết', 'nói',
    'gặp', 'gặp', 'gỡ', 'thảo', 'luận', 'thảo', 'luận', 'bàn', 'bạc', 'quyết', 'định',
    
    # Common adjectives
    'tốt', 'xấu', 'đẹp', 'xấu', 'lớn', 'nhỏ', 'cao', 'thấp', 'dài', 'ngắn',
    'rộng', 'hẹp', 'nhanh', 'chậm', 'nóng', 'lạnh', 'ấm', 'mát', 'sáng', 'tối',
    'mới', 'cũ', 'trẻ', 'già', 'khỏe', 'yếu', 'mạnh', 'yếu', 'thông', 'minh',
    
    # Time and date words
    'năm', 'tháng', 'ngày', 'tuần', 'giờ', 'phút', 'giây', 'sáng', 'trưa', 'chiều',
    'tối', 'đêm', 'hôm', 'nay', 'hôm', 'qua', 'ngày', 'mai', 'tuần', 'sau', 'tháng', 'sau',
    'năm', 'ngoái', 'năm', 'sau', 'thời', 'gian', 'lúc', 'khi', 'trước', 'sau',
    
    # Family and relationships
    'gia', 'đình', 'bố', 'mẹ', 'cha', 'mẹ', 'ông', 'bà', 'anh', 'chị', 'em', 'con',
    'cháu', 'chắt', 'cô', 'chú', 'bác', 'cậu', 'mợ', 'dì', 'cô', 'dượng', 'mẹ', 'kế',
    'bố', 'dượng', 'vợ', 'chồng', 'bạn', 'bè', 'bạn', 'thân', 'người', 'yêu',
    
    # Food and cooking
    'thức', 'ăn', 'món', 'ăn', 'cơm', 'cháo', 'phở', 'bún', 'miến', 'bánh', 'mì',
    'thịt', 'cá', 'tôm', 'cua', 'rau', 'củ', 'quả', 'trái', 'cây', 'hoa', 'quả',
    'nước', 'uống', 'trà', 'cà', 'phê', 'bia', 'rượu', 'nước', 'ngọt', 'sữa',
    
    # Colors and materials
    'màu', 'đỏ', 'xanh', 'vàng', 'trắng', 'đen', 'nâu', 'tím', 'hồng', 'cam',
    'vàng', 'xanh', 'lá', 'xanh', 'dương', 'xám', 'bạc', 'vàng', 'đồng', 'sắt',
    'gỗ', 'đá', 'nhựa', 'vải', 'lụa', 'len', 'cotton', 'polyester', 'nylon',
    
    # Numbers and quantities
    'một', 'hai', 'ba', 'bốn', 'năm', 'sáu', 'bảy', 'tám', 'chín', 'mười',
    'trăm', 'nghìn', 'triệu', 'tỷ', 'tổng', 'cộng', 'trừ', 'nhân', 'chia', 'bằng',
    'nhiều', 'ít', 'tất', 'cả', 'một', 'số', 'vài', 'một', 'vài', 'hầu', 'hết'
]


def preprocess_image_fast(image_path: str) -> List[Image.Image]:
	"""
	Fast image preprocessing for better OCR speed
	"""
	processed_images = []
	
	# Load original image
	img = Image.open(image_path)
	processed_images.append(img)
	
	# Convert to OpenCV format
	img_cv = cv2.imread(image_path)
	
	# 1. Grayscale conversion (fastest)
	gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
	processed_images.append(Image.fromarray(gray))
	
	# 2. Simple contrast enhancement (faster than CLAHE)
	enhanced = cv2.convertScaleAbs(gray, alpha=1.2, beta=10)
	processed_images.append(Image.fromarray(enhanced))
	
	# 3. Single binarization method (fastest effective)
	_, binary = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
	processed_images.append(Image.fromarray(binary))
	
	return processed_images

def preprocess_image(image_path: str) -> List[Image.Image]:
	"""
	Preprocess image to improve OCR accuracy (full version)
	"""
	processed_images = []
	
	# Load original image
	img = Image.open(image_path)
	processed_images.append(img)
	
	# Convert to OpenCV format
	img_cv = cv2.imread(image_path)
	
	# 1. Grayscale conversion
	gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
	processed_images.append(Image.fromarray(gray))
	
	# 2. Noise reduction
	denoised = cv2.medianBlur(gray, 3)
	processed_images.append(Image.fromarray(denoised))
	
	# 3. Contrast enhancement
	clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
	enhanced = clahe.apply(denoised)
	processed_images.append(Image.fromarray(enhanced))
	
	# 4. Binarization with different methods
	# Otsu's method
	_, binary_otsu = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
	processed_images.append(Image.fromarray(binary_otsu))
	
	# Adaptive threshold
	binary_adaptive = cv2.adaptiveThreshold(enhanced, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
	processed_images.append(Image.fromarray(binary_adaptive))
	
	# 5. Morphological operations
	kernel = np.ones((2,2), np.uint8)
	morphed = cv2.morphologyEx(binary_adaptive, cv2.MORPH_CLOSE, kernel)
	processed_images.append(Image.fromarray(morphed))
	
	# 6. Deskewing
	deskewed = deskew_image(morphed)
	processed_images.append(Image.fromarray(deskewed))
	
	return processed_images

def deskew_image(image: np.ndarray) -> np.ndarray:
	"""
	Deskew image to correct rotation
	"""
	# Find contours
	contours, _ = cv2.findContours(image, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
	
	if not contours:
		return image
	
	# Get the largest contour
	largest_contour = max(contours, key=cv2.contourArea)
	
	# Get minimum area rectangle
	rect = cv2.minAreaRect(largest_contour)
	angle = rect[2]
	
	# Correct angle
	if angle < -45:
		angle = 90 + angle
	
	# Rotate image
	(h, w) = image.shape[:2]
	center = (w // 2, h // 2)
	M = cv2.getRotationMatrix2D(center, angle, 1.0)
	rotated = cv2.warpAffine(image, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
	
	return rotated

def correct_vietnamese_text_preserve_structure(text: str) -> str:
	"""
	Correct common OCR mistakes in Vietnamese text while preserving line breaks and structure
	"""
	if not text or not text.strip():
		return text
	
	lines = text.split('\n')
	corrected_lines = []
	
	for line in lines:
		if not line.strip():
			corrected_lines.append(line)
			continue
			
		# Apply corrections
		corrected_line = line
		
		# Fix number recognition errors first
		corrected_line = fix_number_recognition_errors(corrected_line)
		
		# Fix Vietnamese word errors
		corrected_line = fix_vietnamese_word_errors(corrected_line)
		
		# Then apply basic character fixes for Vietnamese text
		basic_fixes = {
			'rn': 'm', 'cl': 'd', 'ci': 'a',
			'vv': 'w', 'ii': 'n', 'nn': 'm'
		}
		
		for mistake, correction in basic_fixes.items():
			corrected_line = corrected_line.replace(mistake, correction)
		
		corrected_lines.append(corrected_line)
	
	return '\n'.join(corrected_lines)

def fix_number_recognition_errors(text: str) -> str:
	"""
	Fix common OCR errors in number recognition.
	"""
	if not text:
		return text
	
	# Common OCR mistakes for numbers
	number_fixes = {
		# Letter 'l' mistaken for '1'
		'l0': '10', 'l1': '11', 'l2': '12', 'l3': '13', 'l4': '14',
		'l5': '15', 'l6': '16', 'l7': '17', 'l8': '18', 'l9': '19',
		'0l': '01', '1l': '11', '2l': '21', '3l': '31', '4l': '41',
		'5l': '51', '6l': '61', '7l': '71', '8l': '81', '9l': '91',
		
		# Letter 'O' mistaken for '0'
		'O0': '00', 'O1': '01', 'O2': '02', 'O3': '03', 'O4': '04',
		'O5': '05', 'O6': '06', 'O7': '07', 'O8': '08', 'O9': '09',
		'0O': '00', '1O': '10', '2O': '20', '3O': '30', '4O': '40',
		'5O': '50', '6O': '60', '7O': '70', '8O': '80', '9O': '90',
		
		# Letter 'S' mistaken for '5'
		'S0': '50', 'S1': '51', 'S2': '52', 'S3': '53', 'S4': '54',
		'S6': '56', 'S7': '57', 'S8': '58', 'S9': '59',
		'0S': '05', '1S': '15', '2S': '25', '3S': '35', '4S': '45',
		'6S': '65', '7S': '75', '8S': '85', '9S': '95',
		
		# Letter 'B' mistaken for '8'
		'B0': '80', 'B1': '81', 'B2': '82', 'B3': '83', 'B4': '84',
		'B5': '85', 'B6': '86', 'B7': '87', 'B9': '89',
		'0B': '08', '1B': '18', '2B': '28', '3B': '38', '4B': '48',
		'5B': '58', '6B': '68', '7B': '78', '9B': '98',
		
		# Letter 'G' mistaken for '6'
		'G0': '60', 'G1': '61', 'G2': '62', 'G3': '63', 'G4': '64',
		'G5': '65', 'G7': '67', 'G8': '68', 'G9': '69',
		'0G': '06', '1G': '16', '2G': '26', '3G': '36', '4G': '46',
		'5G': '56', '7G': '76', '8G': '86', '9G': '96',
		
		# Letter 'I' mistaken for '1'
		'I0': '10', 'I1': '11', 'I2': '12', 'I3': '13', 'I4': '14',
		'I5': '15', 'I6': '16', 'I7': '17', 'I8': '18', 'I9': '19',
		'0I': '01', '1I': '11', '2I': '21', '3I': '31', '4I': '41',
		'5I': '51', '6I': '61', '7I': '71', '8I': '81', '9I': '91',
		
		# Letter 'Z' mistaken for '2'
		'Z0': '20', 'Z1': '21', 'Z3': '23', 'Z4': '24', 'Z5': '25',
		'Z6': '26', 'Z7': '27', 'Z8': '28', 'Z9': '29',
		'0Z': '02', '1Z': '12', '3Z': '32', '4Z': '42', '5Z': '52',
		'6Z': '62', '7Z': '72', '8Z': '82', '9Z': '92',
		
		# Common patterns
		'l9sl': '1951',  # Your specific example
		'l9s': '195',
		'l9': '19',
		'9sl': '951',
	}
	
	# Apply number fixes
	for wrong, correct in number_fixes.items():
		text = text.replace(wrong, correct)
	
	# Fix standalone letters that should be numbers
	# Only replace if they're in a context that suggests they should be numbers
	import re
	
	# Pattern: letter followed by numbers (like l123 -> 1123)
	text = re.sub(r'\b[lOISBZG](?=\d)', lambda m: {
		'l': '1', 'O': '0', 'I': '1', 'S': '5', 'B': '8', 'G': '6', 'Z': '2'
	}.get(m.group(0), m.group(0)), text)
	
	# Pattern: numbers followed by letter (like 123l -> 1231)
	text = re.sub(r'(?<=\d)[lOISBZG]\b', lambda m: {
		'l': '1', 'O': '0', 'I': '1', 'S': '5', 'B': '8', 'G': '6', 'Z': '2'
	}.get(m.group(0), m.group(0)), text)
	
	return text

def fix_vietnamese_word_errors(text: str) -> str:
	"""
	Fix common Vietnamese word recognition errors using context-aware correction.
	"""
	if not text:
		return text
	
	# Vietnamese business and common words dictionary
	vietnamese_words = {
		# Business terms
		'doanh': 'doanh', 'kinh doanh': 'kinh doanh', 'thương mại': 'thương mại',
		'công ty': 'công ty', 'doanh nghiệp': 'doanh nghiệp', 'sản xuất': 'sản xuất',
		'kinh tế': 'kinh tế', 'thị trường': 'thị trường', 'đầu tư': 'đầu tư',
		'ngân hàng': 'ngân hàng', 'tài chính': 'tài chính', 'bảo hiểm': 'bảo hiểm',
		'xuất khẩu': 'xuất khẩu', 'nhập khẩu': 'nhập khẩu', 'thuế': 'thuế',
		'giá cả': 'giá cả', 'chi phí': 'chi phí', 'lợi nhuận': 'lợi nhuận',
		
		# Common words
		'chính phủ': 'chính phủ', 'nhà nước': 'nhà nước', 'xã hội': 'xã hội',
		'giáo dục': 'giáo dục', 'y tế': 'y tế', 'khoa học': 'khoa học',
		'công nghệ': 'công nghệ', 'thông tin': 'thông tin', 'truyền thông': 'truyền thông',
		'văn hóa': 'văn hóa', 'nghệ thuật': 'nghệ thuật', 'thể thao': 'thể thao',
		'du lịch': 'du lịch', 'môi trường': 'môi trường', 'năng lượng': 'năng lượng',
		'giao thông': 'giao thông', 'xây dựng': 'xây dựng', 'nông nghiệp': 'nông nghiệp',
		
		# Technical terms
		'máy tính': 'máy tính', 'phần mềm': 'phần mềm', 'hệ thống': 'hệ thống',
		'dữ liệu': 'dữ liệu', 'mạng': 'mạng', 'internet': 'internet',
		'website': 'website', 'ứng dụng': 'ứng dụng', 'thiết bị': 'thiết bị',
		'điện tử': 'điện tử', 'tự động': 'tự động', 'robot': 'robot',
	}
	
	# Common OCR mistakes for Vietnamese words
	word_fixes = {
		# Common character substitutions
		'đoanh': 'doanh',  # Your specific example
		'kinh đoanh': 'kinh doanh',
		'thương mai': 'thương mại',
		'cong ty': 'công ty',
		'doanh nghiep': 'doanh nghiệp',
		'san xuat': 'sản xuất',
		'kinh te': 'kinh tế',
		'thi truong': 'thị trường',
		'dau tu': 'đầu tư',
		'ngan hang': 'ngân hàng',
		'tai chinh': 'tài chính',
		'bao hiem': 'bảo hiểm',
		'xuat khau': 'xuất khẩu',
		'nhap khau': 'nhập khẩu',
		'gia ca': 'giá cả',
		'chi phi': 'chi phí',
		'loi nhuan': 'lợi nhuận',
		'chinh phu': 'chính phủ',
		'nha nuoc': 'nhà nước',
		'xa hoi': 'xã hội',
		'giao duc': 'giáo dục',
		'khoa hoc': 'khoa học',
		'cong nghe': 'công nghệ',
		'thong tin': 'thông tin',
		'truyen thong': 'truyền thông',
		'van hoa': 'văn hóa',
		'nghe thuat': 'nghệ thuật',
		'the thao': 'thể thao',
		'moi truong': 'môi trường',
		'nang luong': 'năng lượng',
		'giao thong': 'giao thông',
		'xay dung': 'xây dựng',
		'nong nghiep': 'nông nghiệp',
		'may tinh': 'máy tính',
		'phan mem': 'phần mềm',
		'he thong': 'hệ thống',
		'du lieu': 'dữ liệu',
		'ung dung': 'ứng dụng',
		'thiet bi': 'thiết bị',
		'dien tu': 'điện tử',
		'tu dong': 'tự động',
	}
	
	# Apply word fixes
	for wrong, correct in word_fixes.items():
		text = text.replace(wrong, correct)
	
	# Context-aware correction using Levenshtein distance
	import re
	words = text.split()
	corrected_words = []
	
	for word in words:
		# Clean word (remove punctuation)
		clean_word = re.sub(r'[^\w]', '', word.lower())
		
		if len(clean_word) > 2:  # Only process words longer than 2 characters
			# Find best match in dictionary
			best_match = find_best_vietnamese_word_match(clean_word, vietnamese_words)
			if best_match and best_match != clean_word:
				# Preserve original case and punctuation
				corrected_word = word.replace(clean_word, best_match)
				corrected_words.append(corrected_word)
				logger.info(f"Corrected word: '{word}' -> '{corrected_word}'")
			else:
				corrected_words.append(word)
		else:
			corrected_words.append(word)
	
	return ' '.join(corrected_words)

def find_best_vietnamese_word_match(word: str, dictionary: dict) -> str:
	"""
	Find the best match for a word in the Vietnamese dictionary using Levenshtein distance.
	"""
	if not word or len(word) < 3:
		return word
	
	best_match = word
	best_distance = float('inf')
	
	for dict_word in dictionary.keys():
		distance = levenshtein_distance(word, dict_word)
		
		# Only consider matches with reasonable distance
		if distance < best_distance and distance <= max(1, len(word) // 3):
			best_distance = distance
			best_match = dict_word
	
	# Only return correction if distance is reasonable
	if best_distance <= max(1, len(word) // 3):
		return best_match
	
	return word

def correct_vietnamese_text(text: str) -> str:
	"""
	Correct common OCR mistakes in Vietnamese text
	"""
	corrected_text = text
	
	# Apply character corrections
	for mistake, correction in VIETNAMESE_CORRECTIONS.items():
		corrected_text = corrected_text.replace(mistake, correction)
	
	# Fix common word patterns
	patterns = [
		(r'\b(\w+)\s+(\w+)\b', r'\1\2'),  # Remove spaces in words
		(r'(\w+)\s*[|]\s*(\w+)', r'\1l\2'),  # Fix | to l
		(r'(\w+)\s*[0]\s*(\w+)', r'\1o\2'),  # Fix 0 to o
		(r'(\w+)\s*[1]\s*(\w+)', r'\1l\2'),  # Fix 1 to l
		(r'(\w+)\s*[5]\s*(\w+)', r'\1s\2'),  # Fix 5 to s
		(r'(\w+)\s*[8]\s*(\w+)', r'\1B\2'),  # Fix 8 to B
		(r'(\w+)\s*[rn]\s*(\w+)', r'\1m\2'),  # Fix rn to m
		(r'(\w+)\s*[cl]\s*(\w+)', r'\1d\2'),  # Fix cl to d
		(r'(\w+)\s*[ci]\s*(\w+)', r'\1a\2'),  # Fix ci to a
	]
	
	for pattern, replacement in patterns:
		corrected_text = re.sub(pattern, replacement, corrected_text)
	
	# Fix common Vietnamese OCR mistakes
	vietnamese_fixes = {
		# Common OCR mistakes
		'rn': 'm', 'cl': 'd', 'ci': 'a', 'rn': 'm',
		'0': 'o', '1': 'l', '5': 's', '8': 'B',
		'I': 'l', 'O': 'o', 'S': 's',
		# Vietnamese specific
		'rn': 'm', 'cl': 'd', 'ci': 'a',
		# Common word fixes
		'rn': 'm', 'cl': 'd', 'ci': 'a',
		'rn': 'm', 'cl': 'd', 'ci': 'a',
	}
	
	for mistake, correction in vietnamese_fixes.items():
		corrected_text = corrected_text.replace(mistake, correction)
	
	# Context-based correction using Vietnamese words
	words = corrected_text.split()
	corrected_words = []
	
	for word in words:
		# Clean word
		clean_word = re.sub(r'[^\w]', '', word.lower())
		
		# Check if it's similar to a Vietnamese word
		best_match = find_best_vietnamese_match(clean_word)
		if best_match:
			corrected_words.append(best_match)
		else:
			corrected_words.append(word)
	
	return ' '.join(corrected_words)

def find_best_vietnamese_match(word: str) -> Optional[str]:
	"""
	Find the best Vietnamese word match using edit distance
	"""
	if len(word) < 2:
		return None
	
	best_match = None
	best_score = float('inf')
	
	for viet_word in VIETNAMESE_WORDS:
		score = levenshtein_distance(word, viet_word)
		if score < best_score and score <= 2:  # Allow 2 character differences
			best_score = score
			best_match = viet_word
	
	return best_match

def levenshtein_distance(s1: str, s2: str) -> int:
	"""
	Calculate Levenshtein distance between two strings
	"""
	if len(s1) < len(s2):
		return levenshtein_distance(s2, s1)
	
	if len(s2) == 0:
		return len(s1)
	
	previous_row = list(range(len(s2) + 1))
	for i, c1 in enumerate(s1):
		current_row = [i + 1]
		for j, c2 in enumerate(s2):
			insertions = previous_row[j + 1] + 1
			deletions = current_row[j] + 1
			substitutions = previous_row[j] + (c1 != c2)
			current_row.append(min(insertions, deletions, substitutions))
		previous_row = current_row
	
	return previous_row[-1]

def extract_text_from_image(image_path: str, lang: str = "vie+eng", fast_mode: bool = True) -> str:
	"""
	Simple and reliable image text extraction
	"""
	logger.info(f"Processing image: {image_path}")
	
	try:
		# Use original image with simple preprocessing
		img = Image.open(image_path)
		
		# Convert to RGB if needed
		if img.mode != 'RGB':
			img = img.convert('RGB')
		
		# Try with best PSM configs
		configs = [
			'--psm 6',  # Single block
			'--psm 3',  # Single column
			'--psm 8',  # Single word
			'--psm 13'  # Raw line
		]
		best_text = ""
		best_confidence = 0
		
		for config in configs:
			try:
				# Get text with confidence
				data = pytesseract.image_to_data(img, lang=lang, config=config, output_type=pytesseract.Output.DICT)
				
				# Calculate average confidence
				confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
				avg_confidence = sum(confidences) / len(confidences) if confidences else 0
				
				# Extract text
				text = pytesseract.image_to_string(img, lang=lang, config=config)
				
				# If this result is better, keep it
				if avg_confidence > best_confidence and text.strip():
					best_confidence = avg_confidence
					best_text = text
					logger.info(f"Better result found: confidence={avg_confidence:.1f}%, config={config}")
				
			except Exception as e:
				logger.warning(f"OCR failed with config {config}: {e}")
				continue
		
		# Try number-focused OCR if no good result
		if not best_text.strip() or best_confidence < 50:
			logger.info("Trying number-focused OCR...")
			try:
				# Number-focused configuration
				config = '--psm 8 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyzÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚĂĐĨŨƠàáâãèéêìíòóôõùúăđĩũơƯưĂÂÊÔƠĐăâêôơđ'
				text = pytesseract.image_to_string(img, lang=lang, config=config)
				
				if text and len(text.strip()) > len(best_text.strip()):
					best_text = text
					logger.info(f"Number-focused OCR found text: {len(text)} chars")
					
			except Exception as e:
				logger.warning(f"Number-focused OCR failed: {e}")
		
		# If no good result, try with default settings
		if not best_text.strip():
			logger.info("Falling back to default settings")
			best_text = pytesseract.image_to_string(img, lang=lang)
		
		# Apply minimal correction
		if best_text.strip():
			corrected_text = correct_vietnamese_text_preserve_structure(best_text)
			logger.info(f"Text extracted. Length: {len(corrected_text)} characters")
			return corrected_text
		
		return ""
		
	except Exception as e:
		logger.error(f"Failed to process image: {e}")
		return ""


def extract_text_from_word(word_path: str) -> str:
	"""
	Extract text from Word document (.docx) with multiple methods
	"""
	logger.info(f"Processing Word document: {word_path}")
	
	# Method 1: Try docx2txt (simplest and most reliable)
	try:
		logger.info("Trying docx2txt method")
		text = docx2txt.process(word_path)
		if text and text.strip():
			logger.info(f"docx2txt extracted {len(text)} characters")
			logger.info(f"First 200 chars: {text[:200]}")
			# Apply Vietnamese text correction while preserving structure
			corrected_text = correct_vietnamese_text_preserve_structure(text)
			logger.info(f"After correction: {len(corrected_text)} characters")
			return corrected_text
	except Exception as e:
		logger.warning(f"docx2txt method failed: {e}")
	
	# Method 2: Using python-docx library
	try:
		logger.info("Trying python-docx method")
		doc = Document(word_path)
		text_parts = []
		
		logger.info(f"Word document has {len(doc.paragraphs)} paragraphs")
		
		# Extract text from paragraphs
		for i, paragraph in enumerate(doc.paragraphs):
			text = paragraph.text.strip()
			if text:
				text_parts.append(text)
				logger.debug(f"Paragraph {i}: {text[:50]}...")
		
		# Extract text from tables
		logger.info(f"Word document has {len(doc.tables)} tables")
		for table_idx, table in enumerate(doc.tables):
			for row_idx, row in enumerate(table.rows):
				row_text = []
				for cell in row.cells:
					cell_text = cell.text.strip()
					if cell_text:
						row_text.append(cell_text)
				if row_text:
					table_text = ' | '.join(row_text)
					text_parts.append(table_text)
					logger.debug(f"Table {table_idx}, Row {row_idx}: {table_text[:50]}...")
		
		# Extract text from headers and footers
		for section_idx, section in enumerate(doc.sections):
			# Header
			if section.header:
				for para in section.header.paragraphs:
					text = para.text.strip()
					if text:
						text_parts.append(f"[Header] {text}")
			
			# Footer
			if section.footer:
				for para in section.footer.paragraphs:
					text = para.text.strip()
					if text:
						text_parts.append(f"[Footer] {text}")
		
		if text_parts:
			combined_text = '\n'.join(text_parts)
			logger.info(f"python-docx extracted {len(text_parts)} text elements")
			logger.info(f"Total text length: {len(combined_text)} characters")
			logger.info(f"First 200 chars: {combined_text[:200]}")
			
			# Apply Vietnamese text correction while preserving structure
			corrected_text = correct_vietnamese_text_preserve_structure(combined_text)
			logger.info(f"After correction: {len(corrected_text)} characters")
			return corrected_text
		
		logger.warning("No text found in Word document using python-docx")
		
	except Exception as e:
		logger.error(f"python-docx method failed: {e}")
		logger.error(f"Error type: {type(e).__name__}")
		import traceback
		logger.error(f"Traceback: {traceback.format_exc()}")
	
	# Method 3: Fallback ZIP extraction
	try:
		logger.info("Trying fallback ZIP extraction method")
		return extract_text_from_docx_zip(word_path)
	except Exception as e2:
		logger.error(f"All methods failed: {e2}")
		return ""

def extract_text_from_docx_zip(word_path: str) -> str:
	"""
	Fallback method: Extract text from docx as ZIP archive
	"""
	logger.info("Trying fallback method: extracting docx as ZIP")
	
	text_parts = []
	
	try:
		with zipfile.ZipFile(word_path, 'r') as docx_zip:
			logger.info(f"ZIP file contains: {docx_zip.namelist()[:10]}...")
			
			# Read document.xml which contains the main content
			if 'word/document.xml' in docx_zip.namelist():
				document_xml = docx_zip.read('word/document.xml')
				root = ET.fromstring(document_xml)
				
				# Extract text from all text nodes with better handling
				for elem in root.iter():
					if elem.text and elem.text.strip():
						text = elem.text.strip()
						# Skip very short text that might be formatting artifacts
						if len(text) > 1:
							text_parts.append(text)
							logger.debug(f"ZIP extracted: {text[:50]}...")
			
			# Also try to read headers and footers
			for filename in docx_zip.namelist():
				if filename.startswith('word/header') or filename.startswith('word/footer'):
					try:
						content = docx_zip.read(filename)
						root = ET.fromstring(content)
						for elem in root.iter():
							if elem.text and elem.text.strip():
								text = elem.text.strip()
								if len(text) > 1:
									text_parts.append(text)
					except Exception as e:
						logger.debug(f"Error reading {filename}: {e}")
						continue
		
		if text_parts:
			combined_text = '\n'.join(text_parts)
			logger.info(f"ZIP method extracted {len(text_parts)} text elements")
			logger.info(f"Total text length: {len(combined_text)} characters")
			logger.info(f"First 200 chars: {combined_text[:200]}")
			return correct_vietnamese_text_preserve_structure(combined_text)
		
		logger.warning("No text found using ZIP method")
		return ""
		
	except Exception as e:
		logger.error(f"ZIP extraction method failed: {e}")
		import traceback
		logger.error(f"ZIP traceback: {traceback.format_exc()}")
		return ""


def extract_text_from_pdf(pdf_path: str, lang: str = "vie+eng") -> str:
	"""
	Enhanced PDF text extraction with improved OCR
	"""
	logger.info(f"Processing PDF: {pdf_path}")
	
	# 1) Try to read existing text layer first
	text_parts: List[str] = []
	try:
		with pdfplumber.open(pdf_path) as pdf:
			for page_num, page in enumerate(pdf.pages):
				page_text = page.extract_text() or ""
				if page_text.strip():
					text_parts.append(page_text)
					logger.info(f"Extracted text from page {page_num + 1} using text layer")
	except Exception as e:
		logger.warning(f"Failed to extract text layer: {e}")

	if text_parts:
		combined_text = "\n".join(text_parts)
		# Apply text correction while preserving structure
		corrected_text = correct_vietnamese_text_preserve_structure(combined_text)
		return corrected_text

	# 2) If no text layer, convert PDF to images and use enhanced OCR
	logger.info("No text layer found, converting PDF to images for OCR")
	
	try:
		images = convert_from_path(pdf_path, poppler_path=POPPLER_BIN) if POPPLER_BIN else convert_from_path(pdf_path)
		logger.info(f"Converted PDF to {len(images)} images")
		
		for page_num, img in enumerate(images):
			# Save temporary image
			temp_path = f"temp_page_{page_num}.png"
			img.save(temp_path, "PNG")
			
			try:
				# Use fast image processing for better speed
				page_text = extract_text_from_image(temp_path, lang, fast_mode=True)
				if page_text.strip():
					text_parts.append(page_text)
					logger.info(f"Extracted text from page {page_num + 1} using OCR")
			except Exception as e:
				logger.warning(f"Failed to process page {page_num + 1}: {e}")
			finally:
				# Clean up temporary file
				if os.path.exists(temp_path):
					os.remove(temp_path)

		if text_parts:
			combined_text = "\n".join(text_parts)
			# Apply text correction while preserving structure
			corrected_text = correct_vietnamese_text_preserve_structure(combined_text)
			return corrected_text
		else:
			logger.warning("No text could be extracted from PDF")
			return ""
			
	except Exception as e:
		logger.error(f"Failed to process PDF: {e}")
		return ""
